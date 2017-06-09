import docx
from docx import Document
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

companie = "firmamea"
autorizat = "John Doe"

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width

baza =pd.read_csv('table.csv')
k=[]
for i in baza.keys():
    k.append(i)
l=len(baza[k[0]])


doc = Document()

sections=doc.sections
font = doc.styles['Normal'].font
font.name = 'Arial'

for section in sections:
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
for j in range(0,l):
    IDI=baza[k[1]][j]
    tarla=baza[k[0]][j]
    pno=baza[k[2]][j]
    supraf=baza[k[4]][j]
    if pd.isnull(baza[k[5]][j]):
        nrt=''
    else:
        nrt=baza[k[5]][j]
    if pd.isnull(baza[k[6]][j]):
        cf=''
    else:
        cf=int(baza[k[6]][j])
    bol=False
    num=baza[k[3]][j]
    numm=[]
    if num.count('  ')>0:
        bol=True
        n=0
        for i in range(1,len(num)-1):
            if num[i]==' ' and num[i-1]==' ':
                n=i+1
            if num[i]==' ' and num[i+1]==' ' and num[i-1]!=' ':
                numm.append(num[n:i])
                n=i
        numm.append(num[n:len(num)])
    if bol==False:
        numm.append(num)
    lnum=len(numm)

    head =doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h=head.add_run('Anexa nr. 4 - ')
    h_font=h.font
    h_font.size=Pt(10)
    h.bold=True

    hh=head.add_run('    FISA DE DATE A IMOBILULUI')
    hh_font=hh.font
    hh_font.size=Pt(12)
    hh.bold=True

    uatt=doc.add_paragraph()
    #uatt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    uat=uatt.add_run('    UAT BRANISCA')
    uat_font=uat.font
    uat_font=Pt(12)


    sectorr=doc.add_paragraph()
    #sectorr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sector=sectorr.add_run('    Sector cadastral 5')
    sector_font=sector.font
    sector_font=Pt(12)


    IDD=doc.add_paragraph()
    #IDD.alignment= WD_ALIGN_PARAGRAPH.LEFT
    ID=IDD.add_run('    ID imobil %s' % IDI)
    ID_font=ID.font
    ID_font=Pt(12)

    fisaa=doc.add_paragraph()
    fisa=fisaa.add_run('    1. DATE TEREN')
    fisa_font=fisa.font
    fisa_font.size=Pt(12)
    fisa.bold=True

    utable = doc.add_table(rows=3, cols=7)
    utable.alignment = WD_TABLE_ALIGNMENT.CENTER
    utable.style = 'TableGrid'
    urandunu = utable.rows[0].cells
    urandunu[0].text = 'Nr. Tarla/ strada'
    urandunu[1].text = 'Nr. parcela/ nr. postal'
    urandunu[2].text = 'Suprafata masurata'
    urandunu[3].text = 'Nr. CF'
    urandunu[4].paragraphs[0].add_run( 'Nr. cad/')
    urandunu[4].add_paragraph('nr. top')
    urandunu[5].text = 'Imprejmuit/ Neimprejmuit (I/N)'
    urandunu[6].text = 'Zona cooperativizata/ necooperativizata (Co/Nco)'
    uranddoi = utable.rows[1].cells
    uranddoi[0].text = str(tarla)
    uranddoi[1].text = str(pno)
    uranddoi[2].text = str(supraf)
    uranddoi[3].text = str(cf)
    uranddoi[5].text = 'N'
    uranddoi[6].text = 'Co'
    utable.autofit = False
    set_column_width(utable.columns[0], Cm(1.6))
    set_column_width(utable.columns[1], Cm(2.6))
    set_column_width(utable.columns[2], Cm(2.6))
    set_column_width(utable.columns[3], Cm(2.6))
    set_column_width(utable.columns[4], Cm(2.6))
    set_column_width(utable.columns[5], Cm(2.4))
    set_column_width(utable.columns[6], Cm(3.2))

    doc.add_paragraph('    Observatii:')
    datee=doc.add_paragraph()
    date=datee.add_run('    2. DATE CONSTRUCTII PERMANENTE')
    date_font=fisa.font
    date_font.size=Pt(12)
    date.bold=True

    dtable = doc.add_table(rows=2, cols=8)
    dtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    dtable.style = 'TableGrid'
    drandunu = dtable.rows[0].cells
    drandunu[0].text = 'Identificator constructie'
    drandunu[1].text = 'Cod grupa destinatie'
    drandunu[2].text = 'Numar niveluri'
    drandunu[3].paragraphs[0].add_run('Constr. cu acte')
    drandunu[3].add_paragraph('DA/NU')
    drandunu[4].text = 'Constructie condominiu (DA/NU)'
    drandunu[5].text = 'Nr. bloc'
    drandunu[6].text = 'Nr. total UI'
    drandunu[7].text = 'Suprafata construita masurata'
    uranddoi = dtable.rows[1].cells
    uranddoi[0].text = 'C1'
    dtable.autofit = False
    set_column_width(dtable.columns[0], Cm(2.6))
    set_column_width(dtable.columns[1], Cm(2.7))
    set_column_width(dtable.columns[2], Cm(2))
    set_column_width(dtable.columns[3], Cm(2.5))
    set_column_width(dtable.columns[4], Cm(2.7))
    set_column_width(dtable.columns[5], Cm(1.4))
    set_column_width(dtable.columns[6], Cm(1.5))
    set_column_width(dtable.columns[7], Cm(2.25))

    doc.add_paragraph('    Partile comune:')
    datee=doc.add_paragraph()
    date=datee.add_run('    3. PROPRIETATEA / POSESIA')
    date_font=fisa.font
    date_font.size=Pt(12)
    date.bold=True

    ttable = doc.add_table(rows=6, cols=5)
    ttable.alignment = WD_TABLE_ALIGNMENT.CENTER
    ttable.style = 'TableGrid'

    trandunu = ttable.rows[0].cells
    trandunu[0].paragraphs[0].add_run('Nr. Crt.')
    trandunu[1].paragraphs[0].add_run('Nume si prenume detinator/')
    trandunu[1].add_paragraph('Denumire persoana juridica')
    trandunu[2].text = 'CNP/ CUI'
    trandunu[3].text = 'Nr. act de proprietate/ posesie'
    trandunu[4].text = 'Observatii'
    if cf=='':
        act=str(nrt)
    else:
        act=str(cf)
    tranddoi = ttable.rows[1].cells
    tranddoi[0].text = '1'
    tranddoi[1].text = numm[0]
    tranddoi[3].text = act

    trandtrei = ttable.rows[2].cells
    trandtrei[0].text = '2'
    if lnum>1:
        trandtrei[1].text = numm[1]
        trandtrei[3].text = act

    trandpatru = ttable.rows[3].cells
    trandpatru[0].text = '3'
    if lnum>2:
        trandpatru[1].text = numm[2]
        trandpatru[3].text = act

    trandcinci = ttable.rows[4].cells
    trandcinci[0].text = '4'
    if lnum>3:
        trandcinci[1].text = numm[3]
        trandcinci[3].text = act

    trandsase = ttable.rows[5].cells
    trandsase[0].text = '5'
    if lnum>4:
        trandsase[1].text = numm[4]
        trandsase[3].text = act
    #for h in range(0,lnum):

    ttable.autofit = False
    set_column_width(ttable.columns[0], Cm(1.5))
    set_column_width(ttable.columns[1], Cm(5.75))
    set_column_width(ttable.columns[2], Cm(5))
    set_column_width(ttable.columns[3], Cm(3.25))
    set_column_width(ttable.columns[4], Cm(2.25))

    doc.add_paragraph()
    doc.add_paragraph('    Declaratia titularului dreptului de proprietate:')
    doc.add_paragraph('    Subsemnatul......................... domiciliat in .................................... posesor al CI seria..... nr............\
                       eliberat de ........... la data ............... declar ca sunt de acord cu inregistrarea in cartea funciara a dreptului\
                        de proprietate asupra imobilului cu ID nr. ............ reprezentand  ........ cu suprafata diminuata de ......... ha \
                        si cu amplasamentul stabilit conform intelegerii dintre proprietarii imobilelor din sectorul cadastral nr. ........')


    stable = doc.add_table(rows=2, cols=2)
    stable.alignment = WD_TABLE_ALIGNMENT.CENTER
    srandunu = stable.rows[0].cells
    srandunu[0].paragraphs[0].add_run('Posesor/Titular drept de proprietate/ ')
    srandunu[0].add_paragraph('Persoana interesata')
    srandunu[1].paragraphs[0].add_run('')
    srandunu[1].add_paragraph('Reprezentantul Prestatorului')

    sranddoi = stable.rows[1].cells
    sranddoi[0].text = '(nume si prenume, semnatura)'
    sranddoi[1].paragraphs[0].add_run(companie)
    ex=sranddoi[1].add_paragraph(autorizat)


    stable.autofit = False
    set_column_width(stable.columns[0], Cm(12))
    set_column_width(stable.columns[1], Cm(6))
    doc.add_page_break()

doc.save('forms.docx')
